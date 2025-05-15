from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Создание документа
doc = Document()
doc.add_heading("Shooter Game Report", level=1)

doc.add_paragraph("Name: Ivan Ivanov")
doc.add_paragraph("Group: SE-2300")
doc.add_paragraph("Assignment: Shooter Game (Assignment 6)")
doc.add_paragraph("")

doc.add_heading("Project Overview", level=2)
doc.add_paragraph(
    "This project implements a 3D shooter game in Unity based on Assignment 6 guidelines. "
    "Throughout the development, the following features were implemented to fulfill all requirements and additional polish."
)

doc.add_heading("1. Project Setup", level=2)
doc.add_paragraph(
    "• Unity 3D project created\n"
    "• Terrain or plane used as level base\n"
    "• Imported assets for player, enemy, UI, FX\n"
    "[Insert screenshot of Unity setup]"
)

doc.add_heading("2. Player Character", level=2)
doc.add_paragraph(
    "• Player movement: WASD + Shift (run) + Space (jump)\n"
    "• Camera switch: First/Third person via V key\n"
    "• Animator with Idle, Walk, Run, Jump, Die\n"
    "• Health system (100 HP), death animation\n"
    "[Insert screenshot of animator and movement]"
)

doc.add_heading("3. Shooting Mechanics", level=2)
doc.add_paragraph(
    "• Raycast shooting: single and burst fire modes\n"
    "• Ammo system (max 50), reloading (R key)\n"
    "• Muzzle flash, hit particles, sounds\n"
    "• No fire if ammo = 0, click sound plays\n"
    "• Example code:\n"
    "if (Physics.Raycast(ray, out RaycastHit hit, 500f)) {\n"
    "    if (hit.collider.CompareTag(\"Enemy\")) {\n"
    "        hit.collider.GetComponent<EnemyAI>().TakeHit();\n"
    "    }\n"
    "}\n"
    "[Insert screenshot of shooting and ammo UI]"
)

doc.add_heading("4. Enemy AI and Spawning", level=2)
doc.add_paragraph(
    "• 4 spawn points around map\n"
    "• Spawn rate increases every 5 seconds\n"
    "• Enemies chase player, attack in melee\n"
    "• Each hit reduces player HP by 10\n"
    "• 3 hits to kill enemy, with random death animation\n"
    "[Insert screenshot of enemy AI and combat]"
)

doc.add_heading("5. UI & Feedback", level=2)
doc.add_paragraph(
    "• Ammo and health displayed via TextMeshPro\n"
    "• Timer displays how long player survived\n"
    "• Kill streak sound system: 3, 4, 5, 10 kills\n"
    "[Insert screenshot of HUD]"
)

doc.add_heading("6. Game Flow: Death and Victory", level=2)
doc.add_paragraph(
    "• On death: fade to black, red 'YOU ARE DEAD' text, return to menu\n"
    "• On survival: 60 seconds → red 'YOU SURVIVED, WARRIOR', fade to black\n"
    "[Insert screenshot of death/victory screens]"
)

doc.add_heading("7. Scene Flow and Cutscene", level=2)
doc.add_paragraph(
    "• Menu Scene → Video Scene → CadScene → GameScene\n"
    "• VideoPlayer used in VideoScene for intro video\n"
    "• Music stops with fade before video\n"
    "• Scene transition via black fade panels\n"
    "[Insert screenshot of menu and video]"
)

doc.add_heading("8. Extra Features", level=2)
doc.add_paragraph(
    "• Fade-in/out audio transitions\n"
    "• Hit particles stick to enemy\n"
    "• Prop support on level\n"
    "• Proper scene flow and input handling\n"
    "[Insert any extra screenshots]"
)

doc.add_heading("Conclusion", level=2)
doc.add_paragraph(
    "This shooter project was built from scratch using Unity. It includes advanced input handling, "
    "animations, sound, effects, scene control, and enemy AI. The assignment is fully implemented with extra polish."
)

# Сохранение
file_path = "C:\Users/Shooter_Assignment6_Report_IvanIvanov.docx"
doc.save(file_path)
file_path
